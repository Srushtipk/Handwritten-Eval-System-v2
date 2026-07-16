import sys
import os
import docx
import re

def parse_academic_format(doc):
    """Extracts questions and answers from a professor's raw DOCX using regex."""
    questions = []
    current_question = None
    
    # Regex to match 'Q1 a)' or '1 a)' or 'Q 2)' etc.
    q_start_pattern = re.compile(r'^(?:Q\s*)?(\d+\s*[a-z]?\))\s*(.*)', re.IGNORECASE)
    marks_pattern = re.compile(r'marks\s*allotted:\s*(\d+)', re.IGNORECASE)
    
    for para in doc.paragraphs:
        text_block = para.text.strip()
        if not text_block: continue
        
        for text in text_block.split('\n'):
            text = text.strip()
            if not text: continue
            
            # Check if this line starts a new question
            q_match = q_start_pattern.match(text)
            if q_match and len(text) > 10:  # Avoid matching short labels
                # Save the previous question if valid
                if current_question and 'answer' in current_question and current_question['answer'].strip():
                    questions.append(current_question)
                    
                q_id = q_match.group(1).strip()
                q_text = q_match.group(2).strip()
                current_question = {
                    'id': q_id,
                    'question': q_text,
                    'max_marks': 10,  # Default
                    'type': 'flexible',
                    'answer': ''
                }
                continue
            
            if current_question:
                # Look for marks
                m_match = marks_pattern.search(text)
                if m_match:
                    current_question['max_marks'] = int(m_match.group(1))
                    continue 
                    
                # Everything else gets added to the answer/rubric body
                current_question['answer'] += text + "\n"
                
    # Save the last question
    if current_question and 'answer' in current_question and current_question['answer'].strip():
        questions.append(current_question)
        
    return questions

def convert_scheme(input_path, output_path):
    print(f"Reading professor's scheme: {input_path}")
    doc_in = docx.Document(input_path)
    
    questions = parse_academic_format(doc_in)
    
    if not questions:
        print("Error: Could not extract any questions from the document.")
        print("Please ensure questions start with 'Q1 a)' or similar and have 'Marks Allotted: X'.")
        return

    print(f"Successfully extracted {len(questions)} questions. Generating new scheme...")
    
    doc_out = docx.Document()
    
    # Add a title
    doc_out.add_heading('Standardized Marking Scheme', 0)
    
    for i, q in enumerate(questions):
        # 1. Question line
        p_q = doc_out.add_paragraph()
        p_q.add_run(f"Question {q['id']}: ").bold = True
        p_q.add_run(q['question'])
        
        # 2. Max Marks line
        p_m = doc_out.add_paragraph()
        p_m.add_run("Max Marks: ").bold = True
        p_m.add_run(str(q['max_marks']))
        
        # 3. Type line
        p_t = doc_out.add_paragraph()
        p_t.add_run("Type: ").bold = True
        p_t.add_run("flexible")
        
        # 4. Answer line
        p_a = doc_out.add_paragraph()
        p_a.add_run("Answer: ").bold = True
        
        # Add the multi-line answer body
        answer_text = q['answer'].strip()
        if answer_text:
            doc_out.add_paragraph(answer_text)
            
        doc_out.add_paragraph("-" * 40) # Separator
        
    doc_out.save(output_path)
    print(f"\nSUCCESS! Standardized scheme saved to: {output_path}")
    print("You can now safely upload this standardized file to the Web Application.")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python convert_rubric.py <input_prof_scheme.docx> <output_standard_scheme.docx>")
        sys.exit(1)
        
    in_file = sys.argv[1]
    out_file = sys.argv[2]
    
    if not os.path.exists(in_file):
        print(f"Error: Input file '{in_file}' does not exist.")
        sys.exit(1)
        
    convert_scheme(in_file, out_file)
