import docx
import re
import os
import json
from dotenv import load_dotenv
from google import genai

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY", "").strip()

class SchemeParser:
    def __init__(self):
        pass
        
    def parse_scheme(self, docx_path):
        """
        Parses a DOCX marking scheme and returns a list of question dictionaries.
        Expected format:
        Question: ...
        Max Marks: ...
        Type: Flexible/Exact
        Min Length: ... (optional)
        Answer: ...
        """
        doc = docx.Document(docx_path)
        
        questions = []
        current_question = {}
        
        # We will attempt two parsing strategies.
        # Strategy 1: The standard format (Question: / Max Marks: / Answer:)
        questions = self._parse_standard_format(doc)
        
        # Strategy 2: The Academic format (Q1 a) ... / Marks Allotted: ... )
        if not questions:
            questions = self._parse_academic_format(doc)
            
        # Strategy 3: Tabular format (Q. No | Sub | Text | Marks)
        if not questions:
            questions = self._parse_table_format(doc)
            
        # Strategy 4: The Gemini LLM Fallback (for messy unstructured rubrics)
        if not questions and api_key:
            questions = self._parse_with_gemini(doc)
            
        return questions
        
    def _extract_all_text_blocks(self, doc):
        blocks = []
        for para in doc.paragraphs:
            if para.text.strip():
                blocks.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                # Merge cell texts in a row
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    blocks.append(row_text)
        return blocks
        
    def _parse_table_format(self, doc):
        """Parses academic marking schemes that use tables with distinct columns."""
        questions = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if len(cells) >= 3:
                    q_col = cells[0].lower().replace('q', '').strip()
                    # Check if first column is a question number
                    if re.match(r'^\d+$', q_col):
                        q_id = q_col
                        # Handle Sub-question (like 'a)' or 'b')
                        if re.match(r'^[a-z]\)?$', cells[1].lower()):
                            q_id += cells[1].replace(')', '')
                            text_cell = cells[2]
                            marks_cell = cells[3] if len(cells) > 3 else "10"
                        else:
                            text_cell = cells[1]
                            marks_cell = cells[2]
                            
                        # Extract max marks safely
                        marks = 10
                        m = re.search(r'(\d+)', marks_cell)
                        if m: marks = int(m.group(1))
                        
                        # Separate question title from the rest of the answer text
                        lines = [line.strip() for line in text_cell.split('\n') if line.strip()]
                        q_title = lines[0] if lines else "Question"
                        
                        questions.append({
                            'id': q_id,
                            'question': q_title,
                            'max_marks': marks,
                            'type': 'flexible',
                            'answer': text_cell
                        })
        return questions
        
    def _parse_with_gemini(self, doc):
        """Uses Gemini to intelligently extract questions and answers from a messy DOCX."""
        print("Regex parsers failed. Engaging Gemini LLM fallback for scheme extraction...")
        full_text = "\n".join(self._extract_all_text_blocks(doc))
        
        prompt = (
            "You are an expert exam evaluator parsing a university marking scheme/rubric.\n"
            "Extract all questions, their maximum marks, and the ideal answer from the following text.\n"
            "Return a raw JSON list of objects EXACTLY in this format:\n"
            "[\n"
            "  {\n"
            '    "id": "1a",\n'
            '    "question": "What is ...",\n'
            '    "max_marks": 5,\n'
            '    "type": "flexible",\n'
            '    "answer": "The ideal answer text goes here..."\n'
            "  }\n"
            "]\n"
            "Do NOT wrap the output in markdown or ```json. Just return the raw JSON array.\n\n"
            f"--- MARKING SCHEME TEXT ---\n{full_text}"
        )
        
        import time
        for attempt in range(3):
            try:
                client = genai.Client(api_key=api_key)
                response = client.models.generate_content(
                    model='gemini-2.5-flash',
                    contents=prompt
                )
                
                raw = response.text.strip()
                if raw.startswith("```json"):
                    raw = raw[7:]
                if raw.endswith("```"):
                    raw = raw[:-3]
                    
                questions = json.loads(raw.strip())
                return questions
            except Exception as e:
                print(f"Gemini fallback attempt {attempt+1} failed: {e}")
                if attempt < 2:
                    time.sleep(2)
        return []

    def _parse_standard_format(self, doc):
        questions = []
        current_question = {}
        current_key = None
        current_value = ""
        
        def save_field():
            nonlocal current_key, current_value, current_question
            if current_key:
                val = current_value.strip()
                if current_key == 'question':
                    current_question['question'] = val
                elif current_key == 'max_marks':
                    nums = re.findall(r'\d+', val)
                    if nums: current_question['max_marks'] = int(nums[0])
                elif current_key == 'type':
                    current_question['type'] = val.lower()
                elif current_key == 'answer':
                    current_question['answer'] = val
            
        for text_block in self._extract_all_text_blocks(doc):
            if not text_block: continue
                
            for text in text_block.split('\n'):
                text = text.strip()
                if not text: continue
                    
                q_match = re.match(r'^question(.*?):\s*(.*)', text, re.IGNORECASE)
                if q_match:
                    save_field()
                    if 'question' in current_question and 'answer' in current_question:
                        questions.append(current_question)
                    current_question = {}
                    q_id = q_match.group(1).strip() or str(len(questions) + 1)
                    current_question['id'] = q_id
                    current_key = 'question'
                    current_value = q_match.group(2).strip() + "\n"
                    
                elif text.lower().startswith("max marks:"):
                    save_field()
                    current_key = 'max_marks'
                    current_value = text[len("max marks:"):].strip() + "\n"
                    
                elif text.lower().startswith("type:"):
                    save_field()
                    current_key = 'type'
                    current_value = text[len("type:"):].strip() + "\n"
                    
                elif text.lower().startswith("answer:"):
                    save_field()
                    current_key = 'answer'
                    current_value = text[len("answer:"):].strip() + "\n"
                else:
                    if current_key:
                        current_value += text + "\n"
                        
        save_field()
        if 'question' in current_question and 'answer' in current_question:
            questions.append(current_question)
            
        return questions

    def _parse_academic_format(self, doc):
        """Parses academic rubrics like: Q1 a) <question> \\n Marks Allotted: 5 \\n <answer>"""
        questions = []
        current_question = None
        
        # Regex to match 'Q1 a)' or '1 a)' or 'Q 2)' etc.
        q_start_pattern = re.compile(r'^(?:Q\s*)?(\d+\s*[a-z]?\))\s*(.*)', re.IGNORECASE)
        marks_pattern = re.compile(r'marks\s*allotted:\s*(\d+)', re.IGNORECASE)
        
        for text_block in self._extract_all_text_blocks(doc):
            if not text_block: continue
            
            for text in text_block.split('\n'):
                text = text.strip()
                if not text: continue
                
                # Check if this line starts a new question
                q_match = q_start_pattern.match(text)
                if q_match and len(text) > 10:  # Avoid matching short labels
                    # Save the previous question if valid
                    if current_question and 'answer' in current_question and current_question['answer'].strip():
                        questions.append(current_question)
                        
                    q_id = q_match.group(1).strip()
                    q_text = q_match.group(2).strip()
                    current_question = {
                        'id': q_id,
                        'question': q_text,
                        'max_marks': 10,  # Default
                        'type': 'flexible',
                        'answer': ''
                    }
                    continue
                
                if current_question:
                    # Look for marks
                    m_match = marks_pattern.search(text)
                    if m_match:
                        current_question['max_marks'] = int(m_match.group(1))
                        # Don't add the marks metadata line to the answer
                        continue 
                        
                    # Everything else gets added to the answer/rubric body
                    current_question['answer'] += text + "\n"
                    
        # Save the last question
        if current_question and 'answer' in current_question and current_question['answer'].strip():
            questions.append(current_question)
            
        return questions

if __name__ == "__main__":
    pass
