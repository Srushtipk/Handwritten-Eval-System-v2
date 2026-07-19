import docx
doc = docx.Document('IA2-Scheme.docx')
text = []
for para in doc.paragraphs:
    if para.text.strip():
        text.append(para.text.strip())
for table in doc.tables:
    for row in table.rows:
        row_text = " | ".join([cell.text.strip().replace('\n', ' ') for cell in row.cells if cell.text.strip()])
        if row_text:
            text.append(row_text)
with open('dump.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(text))
