import os
os.environ["TOKENIZERS_PARALLELISM"] = "false"
import json
import time
import html
import traceback
import concurrent.futures
import requests
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename

import sys
# Make sure we can import core from parent dir
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from core.ocr_engine import process_pdf, segment_answers_with_gemini
from core.evaluation_engine import HandwrittenEvaluator
from core.scheme_parser import SchemeParser
from core.database import DatabaseManager
import docx

app = Flask(__name__)

# Config
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "data", "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Initialize local offline evaluation engine
try:
    evaluator = HandwrittenEvaluator()
except Exception as e:
    print(f"Warning: Could not initialize local evaluator: {e}")
    evaluator = None

# Initialize local SQLite DB for analytics
db_manager = DatabaseManager(os.path.join(BASE_DIR, "data", "analytics.db"))

scheme_parser = SchemeParser()

# Global dictionary to track batch jobs for SSE progress updates
import threading
batch_jobs = {}
import uuid

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/api/extract', methods=['POST'])
def extract_text():
    if 'scheme' not in request.files or 'student_pdf' not in request.files:
        return jsonify({'error': 'Missing files'}), 400
        
    scheme_file = request.files['scheme']
    student_file = request.files['student_pdf']
    grading_mode = request.form.get('grading_mode', 'experienced')
    
    if scheme_file.filename == '' or student_file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
        
    import uuid
    uid = uuid.uuid4().hex[:8]
    scheme_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uid}_{secure_filename(scheme_file.filename)}")
    student_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{uid}_{secure_filename(student_file.filename)}")
    
    scheme_file.save(scheme_path)
    student_file.save(student_path)
    
    try:
        # Step 1: Parse Scheme
        print("Parsing Scheme DOCX...")
        questions = scheme_parser.parse_scheme(scheme_path)
        if not questions:
            return jsonify({'error': 'No questions found in the DOCX scheme. Check formatting.'}), 400
            
        # Step 2: Extract OCR from Student PDF
        print("Extracting OCR from PDF...")
        ocr_text = process_pdf(student_path)
        if ocr_text.startswith("ERROR"):
            return jsonify({'error': ocr_text}), 500
            
        return jsonify({
            'success': True,
            'ocr_text': ocr_text,
            'questions': questions,
            'grading_mode': grading_mode
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/standardize', methods=['POST'])
def standardize():
    """Endpoint to clean up a messy professor's rubric into the standard format."""
    if 'scheme' not in request.files:
        return jsonify({'error': 'Missing scheme file'}), 400
        
    scheme_file = request.files['scheme']
    if scheme_file.filename == '':
        return jsonify({'error': 'No selected scheme file'}), 400
        
    filename = secure_filename(scheme_file.filename)
    raw_path = os.path.join(UPLOAD_FOLDER, f"raw_{filename}")
    scheme_file.save(raw_path)
    
    clean_filename = f"std_{filename}"
    clean_path = os.path.join(UPLOAD_FOLDER, clean_filename)
    
    try:
        # Use our ultra-robust SchemeParser (which has the Gemini fallback)
        parser = SchemeParser()
        questions = parser.parse_scheme(raw_path)
        
        if not questions:
            return jsonify({'error': 'Standardization failed. The AI could not extract any questions from the document.'}), 400
            
        # Generate the standardized DOCX
        doc_out = docx.Document()
        doc_out.add_heading('Standardized Marking Scheme', 0)
        
        for q in questions:
            p_q = doc_out.add_paragraph()
            p_q.add_run(f"Question {q['id']}: ").bold = True
            p_q.add_run(q['question'])
            
            p_m = doc_out.add_paragraph()
            p_m.add_run("Max Marks: ").bold = True
            p_m.add_run(str(q['max_marks']))
            
            p_t = doc_out.add_paragraph()
            p_t.add_run("Type: ").bold = True
            p_t.add_run(q.get('type', 'flexible'))
            
            p_a = doc_out.add_paragraph()
            p_a.add_run("Answer: ").bold = True
            
            answer_text = q['answer'].strip()
            if answer_text:
                doc_out.add_paragraph(answer_text)
                
            doc_out.add_paragraph("-" * 40)
            
        doc_out.save(clean_path)
            
        # Return the clean file back to the client
        return send_file(
            clean_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=clean_filename
        )
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': f'Internal server error during standardization: {str(e)}'}), 500

@app.route('/api/evaluate', methods=['POST'])
def evaluate_exam():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No JSON payload provided'}), 400
            
        ocr_text = data.get('ocr_text')
        questions = data.get('questions')
        grading_mode = data.get('grading_mode', 'experienced')
        exam_max_override = data.get('exam_max_marks')
        
        student_id = data.get('student_id', 'Unknown Student')
        exam_id = data.get('exam_id', 'Unknown Exam')
        
        semester = data.get('semester')
        subject = data.get('subject')
        subject_code = data.get('subject_code')
        
        if not ocr_text or not questions:
            return jsonify({'error': 'Missing ocr_text or questions in payload'}), 400
            
        # Step 3: Segment Answers
        print("Segmenting Answers with Gemini...")
        segmented_answers = segment_answers_with_gemini(ocr_text, questions)
        
        # Step 4: Evaluate each question
        results = [None] * len(questions)
        total_score = 0
        total_max = 0
        
        print("Starting Evaluation Pipeline...")
        
        def evaluate_question(args):
            i, q = args
            q_safe = q.get('question', '...').encode('ascii', 'ignore').decode()
            print(f"Evaluating Q{i+1}: {q_safe}")
            
            q_text = q.get('question', '')
            ideal = q.get('answer', '')
            max_m = int(q.get('max_marks', 10))
            ans_type = q.get('type', 'flexible').lower()
            min_length = q.get('min_length', None)
            
            q_key = str(q.get('id', str(i+1))).strip()
            
            # Robust fuzzy matching for LLM hallucinated keys
            student_ans = segmented_answers.get(q_key)
            if not student_ans:
                student_ans = segmented_answers.get(f"Q{q_key}")
            if not student_ans:
                student_ans = segmented_answers.get(q_key.replace("Q", ""))
            if not student_ans:
                for k, v in segmented_answers.items():
                    if str(k).lower().strip('q') == q_key.lower().strip('q'):
                        student_ans = v
                        break
                        
            student_ans = (student_ans or "").strip()
            
            if not student_ans or len(student_ans) < 5:
                return i, {
                    'q_num': i+1,
                    'question': q_text,
                    'max_marks': max_m,
                    'type': ans_type.upper(),
                    'score': 0,
                    'reasoning': 'The student did not attempt this question.',
                    'match': 0,
                    'penalty': 0,
                    'extracted_answer': 'Not Attempted',
                    'needs_review': False
                }, max_m
            
            eval_result = evaluator.evaluate(
                question=q_text,
                ideal_rubric=ideal,
                ocr_text=student_ans,
                max_marks=max_m,
                ans_type=ans_type,
                components=q.get('components', {}),
                min_length=min_length,
                grading_mode=grading_mode
            )
            
            return i, {
                'q_num': i + 1,
                'question': q_text,
                'max_marks': max_m,
                'type': ans_type.upper(),
                'score': eval_result['score'],
                'reasoning': html.escape(eval_result['reasoning']),
                'trace': eval_result.get('trace', []),
                'match': eval_result['match_percentage'],
                'penalty': eval_result['penalty'],
                'extracted_answer': html.escape(student_ans),
                'needs_review': eval_result['match_percentage'] < 0.35
            }, max_m
            
        # Process using ThreadPoolExecutor for massive speedup
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(10, len(questions))) as executor:
            # Submit all questions for evaluation
            futures = [executor.submit(evaluate_question, (i, q)) for i, q in enumerate(questions)]
            
            # Wait for them to complete and aggregate
            for future in concurrent.futures.as_completed(futures):
                idx, res, max_m = future.result()
                results[idx] = res
                total_score += res['score']
                total_max += max_m
            
        # Override total_max if provided via UI and scale score proportionally
        if exam_max_override and str(exam_max_override).strip():
            try:
                override_max = int(str(exam_max_override).strip())
                if total_max > 0:
                    # Scale score proportionally, e.g., if you got 14/15, and override is 10, new score is 9/10
                    total_score = int(round((total_score / total_max) * override_max))
                total_max = override_max
            except ValueError:
                pass
                
        # Save evaluation results to database
        evaluation_id = db_manager.save_evaluation(
            exam_id=exam_id,
            student_id=student_id,
            total_score=total_score,
            total_max=total_max,
            grading_mode=grading_mode,
            results=results,
            semester=semester,
            subject=subject,
            subject_code=subject_code
        )
            
        return jsonify({
            'success': True,
            'evaluation_id': evaluation_id,
            'total_score': total_score,
            'total_max': total_max,
            'ocr_extracted': ocr_text,
            'results': results
        })
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/override_score', methods=['POST'])
def override_score():
    try:
        data = request.json
        evaluation_id = data.get('evaluation_id')
        q_num = data.get('q_num')
        new_score = data.get('new_score')
        
        if evaluation_id is None or q_num is None or new_score is None:
            return jsonify({'error': 'Missing required fields'}), 400
            
        new_total = db_manager.override_score(evaluation_id, q_num, int(new_score))
        return jsonify({'success': True, 'new_total': new_total})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/export_csv', methods=['GET'])
def export_csv():
    import csv, io
    from flask import Response
    exam_id = request.args.get('exam_id')
    try:
        data = db_manager.get_export_data(exam_id)
        if not data:
            # Return an empty CSV with headers if no data
            return Response(
                "Student ID,Exam ID,Total Score,Max Marks,Percentage,Grading Mode,Date\n",
                mimetype="text/csv",
                headers={"Content-disposition": "attachment; filename=gradebook_export.csv"}
            )
            
        all_keys = set()
        for row in data:
            all_keys.update(row.keys())
            
        base_cols = ["Student ID", "Exam ID", "Semester", "Subject", "Subject Code", "Total Score", "Max Marks", "Percentage", "Grading Mode", "Date"]
        q_cols = sorted([k for k in all_keys if k.startswith('Q') and k.endswith('Score')], key=lambda x: int(x[1:].split()[0]))
        fieldnames = base_cols + q_cols
        
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=fieldnames)
        writer.writeheader()
        for row in data:
            writer.writerow(row)
            
        return Response(
            output.getvalue(),
            mimetype="text/csv",
            headers={"Content-disposition": "attachment; filename=gradebook_export.csv"}
        )
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/analytics', methods=['GET'])
def get_analytics():
    exam_id = request.args.get('exam_id')
    try:
        summary = db_manager.get_analytics_summary(exam_id)
        return jsonify(summary)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

@app.route('/api/flagged_reviews', methods=['GET'])
def get_flagged_reviews():
    try:
        reviews = db_manager.get_flagged_reviews()
        return jsonify(reviews)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500

def run_batch_job(job_id, scheme_path, pdf_paths, scheme_filename, grading_mode, exam_max_marks, semester, subject, subject_code):
    try:
        # Parse marking scheme once
        questions = scheme_parser.parse_scheme(scheme_path)
        
        batch_jobs[job_id]['total'] = len(pdf_paths)
        
        def process_single_pdf(pdf_path):
            usn = os.path.splitext(os.path.basename(pdf_path))[0]
            try:
                # 1. OCR Extract
                batch_jobs[job_id]['current_usn'] = usn
                ocr_text = process_pdf(pdf_path)
                
                # 2. Segment
                segmented_answers = segment_answers_with_gemini(ocr_text, questions)
                
                # 3. Evaluate Questions
                results = [None] * len(questions)
                total_score = 0
                total_max = 0
                
                for i, q in enumerate(questions):
                    q_text = q.get('question', '')
                    ideal = q.get('answer', '')
                    max_m = int(q.get('max_marks', 10))
                    ans_type = q.get('type', 'flexible').lower()
                    min_length = q.get('min_length', None)
                    q_key = str(q.get('id', str(i+1))).strip()
                    
                    student_ans = segmented_answers.get(q_key)
                    if not student_ans:
                        student_ans = segmented_answers.get(f"Q{q_key}")
                    if not student_ans:
                        student_ans = segmented_answers.get(q_key.replace("Q", ""))
                    if not student_ans:
                        for k, v in segmented_answers.items():
                            if str(k).lower().strip('q') == q_key.lower().strip('q'):
                                student_ans = v
                                break
                    student_ans = (student_ans or "").strip()
                    
                    if not student_ans or len(student_ans) < 5:
                        results[i] = {
                            'q_num': i+1, 'question': q_text, 'max_marks': max_m,
                            'type': ans_type.upper(), 'score': 0, 'reasoning': 'Not attempted.',
                            'match': 0, 'penalty': 0, 'extracted_answer': 'Not Attempted', 'needs_review': False
                        }
                    else:
                        eval_result = evaluator.evaluate(
                            question=q_text, ideal_rubric=ideal, ocr_text=student_ans,
                            max_marks=max_m, ans_type=ans_type, components=q.get('components', {}),
                            min_length=min_length, grading_mode=grading_mode
                        )
                        results[i] = {
                            'q_num': i+1, 'question': q_text, 'max_marks': max_m,
                            'type': ans_type.upper(), 'score': eval_result['score'],
                            'reasoning': eval_result['reasoning'], 'match': eval_result['match_percentage'],
                            'penalty': eval_result['penalty'], 'extracted_answer': student_ans,
                            'needs_review': eval_result['match_percentage'] < 0.35,
                            'trace': eval_result.get('trace', [])
                        }
                    
                    total_score += results[i]['score']
                    total_max += max_m
                
                if exam_max_marks: total_max = int(exam_max_marks)
                
                # 4. Save to DB
                evaluation_id = db_manager.save_evaluation(
                    exam_id=scheme_filename, student_id=usn, total_score=total_score,
                    total_max=total_max, grading_mode=grading_mode, results=results,
                    semester=semester, subject=subject, subject_code=subject_code
                )
                
                # Update progress
                batch_jobs[job_id]['completed'] += 1
                batch_jobs[job_id]['results'].append({
                    'usn': usn, 'total_score': total_score, 'total_max': total_max, 
                    'evaluation_id': evaluation_id, 'data': {'results': results, 'total_score': total_score, 'total_max': total_max}
                })
            except Exception as e:
                batch_jobs[job_id]['completed'] += 1
                batch_jobs[job_id]['errors'].append({'usn': usn, 'error': str(e)})
                traceback.print_exc()

        # Execute parallel
        with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
            executor.map(process_single_pdf, pdf_paths)
            
        batch_jobs[job_id]['status'] = 'completed'
    except Exception as e:
        batch_jobs[job_id]['status'] = 'failed'
        batch_jobs[job_id]['error'] = str(e)
        traceback.print_exc()

@app.route('/api/evaluate_batch_start', methods=['POST'])
def evaluate_batch_start():
    if 'scheme' not in request.files:
        return jsonify({'error': 'Missing scheme file'}), 400
        
    scheme_file = request.files['scheme']
    pdf_files = request.files.getlist('student_pdfs')
    
    if not pdf_files or len(pdf_files) == 0:
        return jsonify({'error': 'Missing PDF files'}), 400
        
    job_id = str(uuid.uuid4())
    job_dir = os.path.join(app.config['UPLOAD_FOLDER'], job_id)
    os.makedirs(job_dir, exist_ok=True)
    
    scheme_path = os.path.join(job_dir, secure_filename(scheme_file.filename))
    scheme_file.save(scheme_path)
    
    pdf_paths = []
    for pdf in pdf_files:
        pdf_path = os.path.join(job_dir, secure_filename(pdf.filename))
        pdf.save(pdf_path)
        pdf_paths.append(pdf_path)
        
    grading_mode = request.form.get('grading_mode', 'experienced')
    exam_max_marks = request.form.get('exam_max_marks')
    semester = request.form.get('semester')
    subject = request.form.get('subject')
    subject_code = request.form.get('subject_code')
    
    batch_jobs[job_id] = {
        'status': 'running',
        'total': len(pdf_paths),
        'completed': 0,
        'current_usn': '',
        'results': [],
        'errors': []
    }
    
    thread = threading.Thread(
        target=run_batch_job, 
        args=(job_id, scheme_path, pdf_paths, scheme_file.filename, grading_mode, exam_max_marks, semester, subject, subject_code)
    )
    thread.daemon = True
    thread.start()
    
    return jsonify({'success': True, 'job_id': job_id})

@app.route('/api/batch_progress/<job_id>', methods=['GET'])
def batch_progress(job_id):
    from flask import Response
    
    def generate():
        while True:
            if job_id not in batch_jobs:
                yield f"data: {json.dumps({'error': 'Job not found'})}\n\n"
                break
                
            job = batch_jobs[job_id]
            yield f"data: {json.dumps(job)}\n\n"
            
            if job['status'] in ['completed', 'failed']:
                break
                
            time.sleep(1.0)
            
    return Response(generate(), mimetype='text/event-stream')

if __name__ == '__main__':
    # Start web server on port 5000. 
    # debug=False prevents the server from violently restarting mid-request when files change.
    app.run(debug=False, port=5000)
